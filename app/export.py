"""交付物导出(V10):Markdown → Word(.docx) / PDF;库 → Excel(.xlsx)."""
import html as _html
import io
import re

from markdown import markdown as _md_lib


def _md_html(md: str) -> str:
    try:
        source = re.sub(r"<[^>\n]{1,4000}>", "", md or "")
        rendered = _md_lib(
            _html.escape(source, quote=False),
            extensions=["tables", "fenced_code"],
        )
        # PDF 导出不需要加载正文中的外部资源；去掉 Markdown 图片和链接标签，
        # 保留可见文字即可。
        rendered = re.sub(r"<img\b[^>]*?/?>", "", rendered, flags=re.I | re.S)
        rendered = re.sub(r"<a\b[^>]*>", "", rendered, flags=re.I | re.S)
        rendered = re.sub(r"</a\s*>", "", rendered, flags=re.I)
        return rendered
    except Exception:
        return "<pre>" + _html.escape(md or "") + "</pre>"


def _blocked_url_fetcher(url: str, *args, **kwargs):
    """WeasyPrint 不得读取本机文件、内网或任意外部 URL。"""
    raise ValueError(f"external resource blocked: {str(url)[:80]}")


def md_to_pdf(md: str, title: str = "交付物") -> bytes:
    from weasyprint import HTML
    html = f"""<html><head><meta charset="utf-8"><style>
    @page {{ size: A4; margin: 2cm 1.8cm; }}
    body {{ font-family: "Noto Sans CJK SC","Noto Serif CJK SC",sans-serif;
           font-size: 11pt; color: #222; line-height: 1.75; }}
    h1 {{ font-size: 18pt; border-bottom: 2px solid #e8b93f; padding-bottom: 6px; }}
    h2 {{ font-size: 14pt; margin-top: 18px; }}
    h3 {{ font-size: 12pt; }}
    table {{ border-collapse: collapse; width: 100%; font-size: 9.5pt; }}
    th, td {{ border: 1px solid #999; padding: 5px 8px; }}
    th {{ background: #f6edd3; }}
    blockquote {{ border-left: 3px solid #e8b93f; margin: 8px 0; padding: 4px 12px; color: #666; }}
    footer {{ font-size: 8pt; color: #999; margin-top: 24px; }}
    </style></head><body>{_md_html(md)}
    <footer>由「派活 PaiHuo」数字员工协助生成 · 请在发布前核对事实与业务数据</footer></body></html>"""
    return HTML(string=html, url_fetcher=_blocked_url_fetcher).write_pdf()


def md_to_docx(md: str, title: str = "交付物") -> bytes:
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    lines = (md or "").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s:\-|]+\|?$", lines[i + 1].strip()):
            # 管道表格
            header = [c.strip() for c in line.strip("|").split("|")]
            rows = []
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            t = doc.add_table(rows=1 + len(rows), cols=len(header))
            t.style = "Table Grid"
            for c, h in enumerate(header):
                t.rows[0].cells[c].text = h
            for r, row in enumerate(rows):
                for c in range(len(header)):
                    t.rows[r + 1].cells[c].text = row[c] if c < len(row) else ""
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif re.match(r"^[-*] ", line):
            doc.add_paragraph(re.sub(r"^[-*] ", "", line).replace("**", ""), style="List Bullet")
        elif re.match(r"^\d+\. ", line):
            doc.add_paragraph(re.sub(r"^\d+\. ", "", line).replace("**", ""), style="List Number")
        elif line.startswith("> "):
            doc.add_paragraph(line[2:].replace("**", ""), style="Intense Quote")
        elif line.strip():
            doc.add_paragraph(line.replace("**", ""))
        i += 1
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _xlsx_safe(value) -> str:
    """防 CSV/Excel 公式注入:客资、线索等表格含用户可控文本,以 = + - @ 开头的单元格
    会被 Excel 当公式执行(如 =HYPERLINK(...) 外带数据)。加前导单引号强制文本。"""
    text = str(value if value is not None else "")
    if text[:1] in ("=", "+", "-", "@", "\t", "\r"):
        return "'" + text
    return text


def rows_to_xlsx(rows: list, headers: list, sheet: str = "数据") -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill
    wb = Workbook()
    ws = wb.active
    ws.title = sheet
    ws.append([_xlsx_safe(h) for h in headers])
    for c in ws[1]:
        c.font = Font(bold=True)
        c.fill = PatternFill("solid", fgColor="FFF0D9")
    for r in rows:
        ws.append([_xlsx_safe(r.get(h, "")) for h in headers])
    for col in ws.columns:
        width = max(len(str(c.value or "")) for c in col)
        ws.column_dimensions[col[0].column_letter].width = min(max(width + 2, 8), 50)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()
